import os
import pandas as pd
import tempfile
from docx import Document
from docx.shared import Pt
from openpyxl import load_workbook

class ExcelProcessorError(Exception):
    pass

def is_cell_empty(value) -> bool:
    if value is None:
        return True
    if isinstance(value, str) and value.strip() == "":
        return True
    return False

def get_max_used_column(ws) -> int:
    max_used_col = 0
    max_row = ws.max_row or 0

    for row in ws.iter_rows(min_row=1, max_row=max_row):
        last_non_empty = 0
        for idx, cell in enumerate(row, start=1):
            if not is_cell_empty(cell.value):
                last_non_empty = idx
        if last_non_empty > max_used_col:
            max_used_col = last_non_empty

    return max_used_col

def validate_file_exists(file_path: str) -> None:
    if not os.path.exists(file_path):
        raise ExcelProcessorError(f"File không tồn tại: {file_path}")
    if not os.path.isfile(file_path):
        raise ExcelProcessorError(f"Đường dẫn không phải file: {file_path}")

def validate_excel_file(file_path: str) -> None:
    validate_file_exists(file_path)
    ext = os.path.splitext(file_path)[1].lower()
    
    if ext == '.xls':
         raise ExcelProcessorError("Hệ thống hiện tại chỉ hỗ trợ file .xlsx (OpenXML). Vui lòng lưu file sang định dạng .xlsx và thử lại.")
         
    if ext not in (".xlsx", ".xls"):
        raise ExcelProcessorError(f"File không phải Excel: {ext}")

    if os.path.getsize(file_path) > 50 * 1024 * 1024:
        raise ExcelProcessorError("File quá lớn (tối đa 50MB)")

def get_sheet_names(file_path: str) -> list[str]:
    validate_excel_file(file_path)
    try:
        wb = load_workbook(file_path, read_only=True, data_only=True)
        names = wb.sheetnames
        wb.close()
        return names
    except Exception as e:
        raise ExcelProcessorError(f"Không thể đọc file Excel: {str(e)}")

def preview_sheet_data(file_path: str, sheet_name: str, num_rows: int = 10) -> dict:
    """
    Preview = quét TOÀN BỘ sheet, lấy N dòng CÓ DỮ LIỆU đầu tiên.
    KHÔNG dùng header_row, KHÔNG dùng data_start_row.
    """
    validate_excel_file(file_path)

    wb = load_workbook(file_path, read_only=True, data_only=True)
    if sheet_name not in wb.sheetnames:
        raise ExcelProcessorError(f"Sheet '{sheet_name}' không tồn tại")

    ws = wb[sheet_name]
    max_row = ws.max_row or 0
    max_used_col = get_max_used_column(ws)

    preview = []
    for row in ws.iter_rows(min_row=1, max_row=max_row, max_col=max_used_col):
        values = ["" if c.value is None else c.value for c in row]
        preview.append(values)

    wb.close()

    return {
        "preview": preview,
        "total_rows": max_row,
        "total_cols": max_used_col,
    }

def get_column_headers(file_path: str, sheet_name: str, header_row: int) -> list[str]:
    """
    SỬA LỖI QUAN TRỌNG:
    Dùng iter_rows thay vì truy cập trực tiếp index để tránh lỗi với openpyxl read_only
    """
    validate_excel_file(file_path)

    try:
        wb = load_workbook(file_path, read_only=True, data_only=True)
        if sheet_name not in wb.sheetnames:
            raise ExcelProcessorError(f"Sheet '{sheet_name}' không tồn tại")

        ws = wb[sheet_name]
        
        if ws.max_row and header_row > ws.max_row:
             wb.close()
             raise ExcelProcessorError(f"Dòng header ({header_row}) lớn hơn tổng số dòng ({ws.max_row})")

        headers = []
        found_row = False
        max_used_col = get_max_used_column(ws)

        row_generator = ws.iter_rows(
            min_row=header_row,
            max_row=header_row,
            max_col=max_used_col
        )
        
        try:
            row_cells = next(row_generator)
            found_row = True
            
            for idx, cell in enumerate(row_cells, start=1):
                if len(headers) > 100 and cell.value is None: 
                    continue 

                if cell.value is None or str(cell.value).strip() == "":
                    headers.append(f"Cột {idx}")
                else:
                    headers.append(
                        str(cell.value).replace("\n", " ").replace("\t", " ").strip()
                    )
        except StopIteration:
            pass

        wb.close()
        
        if not found_row or not headers:
             return []

        return headers

    except Exception as e:
        if isinstance(e, ExcelProcessorError):
            raise e
        raise ExcelProcessorError(f"Lỗi khi đọc cột: {str(e)}")

def convert_excel_to_docx(
    excel_file_path: str,
    output_docx_path: str,
    sheet_name: str,
    selected_columns: list[str],
    header_row: int,
    data_start_row: int,
    data_end_row: int | None = None,
) -> int:
    validate_excel_file(excel_file_path)

    if not selected_columns:
        raise ExcelProcessorError("Chưa chọn cột để xuất")

    if data_start_row <= header_row:
        raise ExcelProcessorError("Dòng data phải > dòng header")

    try:
        df = pd.read_excel(
            excel_file_path,
            sheet_name=sheet_name,
            header=header_row - 1,
            dtype=str
        )
    except Exception as e:
         raise ExcelProcessorError(f"Lỗi đọc dữ liệu Excel: {str(e)}")

    # Chuẩn hóa tên cột
    df.columns = (
        df.columns.astype(str)
        .str.replace("\n", " ")
        .str.replace("\t", " ")
        .str.strip()
    )
    
    # Kiểm tra cột
    missing = [c for c in selected_columns if c not in df.columns]
    if missing:
        raise ExcelProcessorError(f"Không tìm thấy các cột sau: {', '.join(missing)}")

    start_idx = data_start_row - header_row - 1
    end_idx = None
    if data_end_row:
        end_idx = data_end_row - header_row

    if start_idx < 0: start_idx = 0

    df_subset = df.iloc[start_idx:end_idx].copy()
    
    df_final = df_subset[selected_columns].reset_index(drop=True)
    
    df_final = df_final.fillna("")

    df_final = df_final.replace(r'^\s*$', pd.NA, regex=True).ffill()
    
    df_final = df_final.fillna("")

    if df_final.empty:
        raise ExcelProcessorError("Không có dữ liệu nào trong khoảng dòng đã chọn")

    try:
        doc = Document()
        style = doc.styles["Normal"]
        style.font.name = "Arial"
        style.font.size = Pt(11)

        total_rows = len(df_final)

        for idx, (_, row) in enumerate(df_final.iterrows()):
            p = doc.add_paragraph()

            for col in selected_columns:
                val = str(row[col]).strip()

                run_header = p.add_run(f"{col}: ")
                run_header.bold = True
                p.add_run(f"{val}\n")

            if idx < total_rows - 1:
                doc.add_paragraph("-" * 50)


        os.makedirs(os.path.dirname(output_docx_path), exist_ok=True)
        doc.save(output_docx_path)

        return total_rows

    except Exception as e:
        raise ExcelProcessorError(f"Lỗi khi ghi file DOCX: {str(e)}")

def convert_excel_to_markdown(
    excel_file_path: str,
    output_md_path: str,
    sheet_name: str,
    selected_columns: list[str],
    header_row: int,
    data_start_row: int,
    data_end_row: int | None = None,
) -> int:
    validate_excel_file(excel_file_path)

    if not selected_columns:
        raise ExcelProcessorError("Chưa chọn cột để xuất")

    if data_start_row <= header_row:
        raise ExcelProcessorError("Dòng data phải > dòng header")

    try:
        df = pd.read_excel(
            excel_file_path,
            sheet_name=sheet_name,
            header=header_row - 1,
            dtype=str
        )
    except Exception as e:
        raise ExcelProcessorError(f"Lỗi đọc dữ liệu Excel: {str(e)}")

    df.columns = (
        df.columns.astype(str)
        .str.replace("\n", " ")
        .str.replace("\t", " ")
        .str.strip()
    )

    missing = [c for c in selected_columns if c not in df.columns]
    if missing:
        raise ExcelProcessorError(f"Không tìm thấy các cột sau: {', '.join(missing)}")

    start_idx = data_start_row - header_row - 1
    end_idx = None
    if data_end_row:
        end_idx = data_end_row - header_row

    if start_idx < 0:
        start_idx = 0

    df_subset = df.iloc[start_idx:end_idx].copy()
    df_final = df_subset[selected_columns].reset_index(drop=True)
    df_final = df_final.fillna("")
    df_final = df_final.replace(r'^\s*$', pd.NA, regex=True).ffill()
    df_final = df_final.fillna("")

    if df_final.empty:
        raise ExcelProcessorError("Không có dữ liệu nào trong khoảng dòng đã chọn")

    temp_path = None
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".xlsx") as temp_file:
            temp_path = temp_file.name

        df_final.to_excel(temp_path, index=False)

        try:
            from markitdown import MarkItDown
        except Exception:
            raise ExcelProcessorError(
                "Chưa cài markitdown. Cài bằng: pip install markitdown[xlsx]"
            )

        md = MarkItDown()
        result = md.convert(temp_path)
        markdown_text = getattr(result, "text_content", None)
        if markdown_text is None:
            markdown_text = str(result)

        os.makedirs(os.path.dirname(output_md_path), exist_ok=True)
        with open(output_md_path, "w", encoding="utf-8") as f:
            f.write(markdown_text)

        return len(df_final)
    except ExcelProcessorError:
        raise
    except Exception as e:
        raise ExcelProcessorError(f"Lỗi khi tạo Markdown: {str(e)}")
    finally:
        if temp_path and os.path.exists(temp_path):
            os.remove(temp_path)
