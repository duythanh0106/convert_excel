import os
import tempfile
from typing import Optional

import pandas as pd
from docx import Document
from docx.shared import Pt
from openpyxl import load_workbook


class ExcelProcessorError(Exception):
    pass


class ClassicExcelProcessor:
    def __init__(self, max_excel_size_mb: int = 50) -> None:
        self.max_excel_size = max_excel_size_mb * 1024 * 1024

    @staticmethod
    def _is_cell_empty(value) -> bool:
        if value is None:
            return True
        return isinstance(value, str) and value.strip() == ""

    def _validate_file_exists(self, file_path: str) -> None:
        if not os.path.exists(file_path):
            raise ExcelProcessorError(f"File không tồn tại: {file_path}")
        if not os.path.isfile(file_path):
            raise ExcelProcessorError(f"Đường dẫn không phải file: {file_path}")

    def _validate_excel_file(self, file_path: str) -> None:
        self._validate_file_exists(file_path)

        ext = os.path.splitext(file_path)[1].lower()
        if ext == ".xls":
            raise ExcelProcessorError(
                "Hệ thống hiện tại chỉ hỗ trợ file .xlsx (OpenXML). "
                "Vui lòng lưu file sang định dạng .xlsx và thử lại."
            )
        if ext not in (".xlsx", ".xls"):
            raise ExcelProcessorError(f"File không phải Excel: {ext}")

        if os.path.getsize(file_path) > self.max_excel_size:
            raise ExcelProcessorError(f"File quá lớn (tối đa {self.max_excel_size // (1024*1024)}MB)")

    def _get_max_used_column(self, ws) -> int:
        max_used_col = 0
        max_row = ws.max_row or 0

        for row in ws.iter_rows(min_row=1, max_row=max_row):
            last_non_empty = 0
            for idx, cell in enumerate(row, start=1):
                if not self._is_cell_empty(cell.value):
                    last_non_empty = idx
            max_used_col = max(max_used_col, last_non_empty)

        return max_used_col

    @staticmethod
    def _norm_header(x) -> str:
        return str(x).replace("\n", " ").replace("\t", " ").strip()

    def _ensure_output_dir(self, output_path: str) -> None:
        out_dir = os.path.dirname(output_path)
        if out_dir:
            os.makedirs(out_dir, exist_ok=True)

    def _read_df(
        self,
        excel_file_path: str,
        sheet_name: str,
        header_row: int,
        dtype=str
    ) -> pd.DataFrame:
        try:
            df = pd.read_excel(
                excel_file_path,
                sheet_name=sheet_name,
                header=header_row - 1,
                dtype=dtype,
            )
        except Exception as e:
            raise ExcelProcessorError(f"Lỗi đọc dữ liệu Excel: {str(e)}")

        df.columns = (
            df.columns.astype(str)
            .str.replace("\n", " ")
            .str.replace("\t", " ")
            .str.strip()
        )
        return df

    def _slice_and_clean(
        self,
        df: pd.DataFrame,
        selected_columns: list[str],
        header_row: int,
        data_start_row: int,
        data_end_row: int | None,
    ) -> pd.DataFrame:
        if not selected_columns:
            raise ExcelProcessorError("Chưa chọn cột để xuất")

        if data_start_row <= header_row:
            raise ExcelProcessorError("Dòng data phải > dòng header")

        missing = [c for c in selected_columns if c not in df.columns]
        if missing:
            raise ExcelProcessorError(f"Không tìm thấy các cột sau: {', '.join(missing)}")

        start_idx = data_start_row - header_row - 1
        if start_idx < 0:
            start_idx = 0

        end_idx = None
        if data_end_row:
            end_idx = data_end_row - header_row

        df_final = df.iloc[start_idx:end_idx][selected_columns].copy().reset_index(drop=True)

        df_final = df_final.fillna("")
        df_final = df_final.replace(r"^\s*$", pd.NA, regex=True).ffill()
        df_final = df_final.fillna("")

        if df_final.empty:
            raise ExcelProcessorError("Không có dữ liệu nào trong khoảng dòng đã chọn")

        return df_final

    def get_sheet_names(self, file_path: str) -> list[str]:
        self._validate_excel_file(file_path)
        try:
            wb = load_workbook(file_path, read_only=True, data_only=True)
            try:
                return wb.sheetnames
            finally:
                wb.close()
        except Exception as e:
            raise ExcelProcessorError(f"Không thể đọc file Excel: {str(e)}")

    def preview_sheet_data(self, file_path: str, sheet_name: str, num_rows: int = 20) -> dict:
        self._validate_excel_file(file_path)

        wb = load_workbook(file_path, read_only=True, data_only=True)
        try:
            if sheet_name not in wb.sheetnames:
                raise ExcelProcessorError(f"Sheet '{sheet_name}' không tồn tại")

            ws = wb[sheet_name]
            max_row = ws.max_row or 0
            max_used_col = self._get_max_used_column(ws)

            if max_used_col == 0:
                return {"preview": [], "total_rows": 0, "total_cols": 0, "display_rows": 0}

            preview = []
            last_data_row = 0
            preview_limit = max(1, min(num_rows, 200))  

            for row_index, row in enumerate(
                ws.iter_rows(min_row=1, max_row=max_row, max_col=max_used_col),
                start=1
            ):
                values = []
                has_data = False
                for cell in row:
                    v = cell.value
                    if isinstance(v, str):
                        v = v.strip()
                    values.append("" if v is None else v)
                    if not self._is_cell_empty(v):
                        has_data = True

                if has_data:
                    last_data_row = row_index

                if len(preview) < preview_limit:
                    preview.append(values)

            if last_data_row < len(preview):
                preview = preview[:last_data_row]

            return {
                "preview": preview,
                "total_rows": last_data_row,
                "total_cols": max_used_col,
                "display_rows": len(preview),
            }
        finally:
            wb.close()

    def get_column_headers(self, file_path: str, sheet_name: str, header_row: int) -> list[str]:
        self._validate_excel_file(file_path)

        def read_headers(data_only: bool) -> list[str]:
            wb = load_workbook(file_path, read_only=True, data_only=data_only)
            try:
                if sheet_name not in wb.sheetnames:
                    raise ExcelProcessorError(f"Sheet '{sheet_name}' không tồn tại")

                ws = wb[sheet_name]
                if ws.max_row and header_row > ws.max_row:
                    raise ExcelProcessorError(
                        f"Dòng header ({header_row}) lớn hơn tổng số dòng ({ws.max_row})"
                    )

                max_used_col = self._get_max_used_column(ws)
                if max_used_col == 0:
                    return []

                row_gen = ws.iter_rows(
                    min_row=header_row,
                    max_row=header_row,
                    max_col=max_used_col
                )

                try:
                    cells = next(row_gen)
                except StopIteration:
                    return []

                headers = []
                for c in cells:
                    if c.value is None:
                        continue
                    h = self._norm_header(c.value)
                    if h:
                        headers.append(h)
                return headers
            finally:
                wb.close()

        try:
            headers = read_headers(data_only=True)
            return headers if headers else read_headers(data_only=False)
        except ExcelProcessorError:
            raise
        except Exception as e:
            raise ExcelProcessorError(f"Lỗi khi đọc cột: {str(e)}")

    def convert_excel_to_docx(
        self,
        excel_file_path: str,
        output_docx_path: str,
        sheet_name: str,
        selected_columns: list[str],
        header_row: int,
        data_start_row: int,
        data_end_row: int | None = None,
    ) -> int:
        self._validate_excel_file(excel_file_path)

        df = self._read_df(excel_file_path, sheet_name, header_row, dtype=str)
        df_final = self._slice_and_clean(df, selected_columns, header_row, data_start_row, data_end_row)

        try:
            doc = Document()
            style = doc.styles["Normal"]
            style.font.name = "Arial"
            style.font.size = Pt(11)

            total_rows = len(df_final)

            for idx, (_, row) in enumerate(df_final.iterrows()):
                p = doc.add_paragraph()
                for col in selected_columns:
                    val = str(row[col]).strip()
                    run_header = p.add_run(f"{col}: ")
                    run_header.bold = True
                    p.add_run(f"{val}\n")

                if idx < total_rows - 1:
                    doc.add_paragraph("-" * 50)

            self._ensure_output_dir(output_docx_path)
            doc.save(output_docx_path)
            return total_rows

        except Exception as e:
            raise ExcelProcessorError(f"Lỗi khi ghi file DOCX: {str(e)}")


class UniversalFileProcessor:
    def __init__(self, max_file_size: int = 100 * 1024 * 1024) -> None:
        self.max_file_size = max_file_size
        self._converter = None
        self._converter_error: Optional[Exception] = None

        try:
            from markitdown import MarkItDown
            self._converter = MarkItDown()
        except Exception as exc:
            self._converter_error = exc

    def _validate_file(self, file_path: str) -> None:
        if not os.path.exists(file_path):
            raise ExcelProcessorError(f"File không tồn tại: {file_path}")
        if not os.path.isfile(file_path):
            raise ExcelProcessorError(f"Đường dẫn không phải file: {file_path}")

        size = os.path.getsize(file_path)
        if size > self.max_file_size:
            raise ExcelProcessorError(
                f"File quá lớn ({size / (1024*1024):.2f}MB). "
                f"Tối đa: {self.max_file_size / (1024*1024):.0f}MB"
            )

    @staticmethod
    def _ensure_output_dir(output_path: str) -> None:
        out_dir = os.path.dirname(output_path)
        if out_dir:
            os.makedirs(out_dir, exist_ok=True)

    def convert_to_markdown(self, file_path: str, output_path: Optional[str] = None) -> str:
        if self._converter is None:
            detail = str(self._converter_error) if self._converter_error else "markitdown"
            raise ExcelProcessorError(f"Chưa cài markitdown: {detail}")

        self._validate_file(file_path)

        try:
            result = self._converter.convert(file_path)
            markdown_text = getattr(result, "text_content", None)
            if markdown_text is None:
                markdown_text = str(result)

            if output_path:
                self._ensure_output_dir(output_path)
                with open(output_path, "w", encoding="utf-8") as f:
                    f.write(markdown_text)

            return markdown_text
        except ExcelProcessorError:
            raise
        except Exception as e:
            raise ExcelProcessorError(f"Lỗi khi convert sang Markdown: {str(e)}")
